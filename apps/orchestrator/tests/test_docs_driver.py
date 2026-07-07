"""docs driver — jailed extraction must stay inside the vault and never guess."""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest

_SERVER = Path(__file__).resolve().parents[1] / "tools" / "docs_server.py"

spec = importlib.util.spec_from_file_location("docs_server", _SERVER)
docs_server = importlib.util.module_from_spec(spec)
spec.loader.exec_module(docs_server)


@pytest.fixture
def vault(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setattr(docs_server, "VAULT", tmp_path.resolve())
    return tmp_path


def test_docx_text_and_tables_extracted(vault: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Quarterly invoice for Acme Corp.")
    doc.add_paragraph("Total due: 4,250.00 EUR by 2026-08-01.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Consulting retainer"
    doc.save(vault / "invoice.docx")

    text = docs_server.read_document("invoice.docx")
    assert "Acme Corp" in text
    assert "4,250.00 EUR" in text
    assert "Consulting retainer" in text


def test_blank_pdf_reports_no_text_instead_of_guessing(vault: Path):
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open(vault / "scan.pdf", "wb") as fh:
        writer.write(fh)

    text = docs_server.read_document("scan.pdf")
    assert "no extractable text" in text


def test_txt_passthrough_and_truncation(vault: Path):
    (vault / "big.txt").write_text("x" * 30_000, encoding="utf-8")
    text = docs_server.read_document("big.txt")
    assert text.endswith("...[truncated]")
    assert len(text) < 30_000


def test_path_escape_is_rejected(vault: Path):
    (vault.parent / "secret.txt").write_text("nope", encoding="utf-8")
    with pytest.raises(ValueError, match="escapes vault"):
        docs_server.read_document("../secret.txt")


def test_unsupported_extension_rejected(vault: Path):
    (vault / "binary.exe").write_bytes(b"MZ")
    with pytest.raises(ValueError, match="Unsupported type"):
        docs_server.read_document("binary.exe")


def test_missing_file_raises(vault: Path):
    with pytest.raises(FileNotFoundError):
        docs_server.read_document("nope.pdf")


def test_list_documents_finds_pdf_and_docx_but_not_system(vault: Path):
    (vault / "00-Inbox").mkdir()
    (vault / "00-Inbox" / "a.pdf").write_bytes(b"%PDF-1.4")
    (vault / "00-Inbox" / "b.docx").write_bytes(b"PK")
    (vault / ".system").mkdir()
    (vault / ".system" / "hidden.pdf").write_bytes(b"%PDF-1.4")

    docs = docs_server.list_documents()
    assert docs == ["00-Inbox/a.pdf", "00-Inbox/b.docx"]
