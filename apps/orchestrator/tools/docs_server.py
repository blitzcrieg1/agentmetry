"""docs — read-only PDF/DOCX text extraction MCP driver, jailed to the vault.

Local-first: pure-Python extraction (pypdf / python-docx), no network, no
writes. Mounted via vault/.system/drivers.json:

    { "name": "docs", "command": "<venv python>",
      "args": ["<this file>", "<vault path>"] }

Drop a PDF or DOCX anywhere in the vault (e.g. 00-Inbox/contract.pdf) and a
skill can pull its text through docs.read_document for summarization.
"""

from __future__ import annotations

import sys
from pathlib import Path

from mcp.server.fastmcp import FastMCP

VAULT = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else Path.cwd() / "vault"

server = FastMCP("docs")

_MAX_CHARS = 24_000
_MAX_PAGES = 50
_SUPPORTED = (".pdf", ".docx", ".txt", ".md")


def _safe(rel: str) -> Path:
    path = (VAULT / rel).resolve()
    if path != VAULT and VAULT not in path.parents:
        raise ValueError(f"Path escapes vault: {rel}")
    return path


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) <= _MAX_CHARS:
        return text
    return text[:_MAX_CHARS] + "\n...[truncated]"


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = reader.pages[:_MAX_PAGES]
    parts = []
    for i, page in enumerate(pages, 1):
        extracted = page.extract_text() or ""
        if extracted.strip():
            parts.append(f"[page {i}]\n{extracted.strip()}")
    if len(reader.pages) > _MAX_PAGES:
        parts.append(f"...[{len(reader.pages) - _MAX_PAGES} more pages omitted]")
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


@server.tool()
def read_document(path: str) -> str:
    """Extract plain text from a vault-relative PDF, DOCX, TXT, or MD file."""
    target = _safe(path)
    if not target.is_file():
        raise FileNotFoundError(f"No such document: {path}")
    suffix = target.suffix.lower()
    if suffix not in _SUPPORTED:
        raise ValueError(f"Unsupported type '{suffix}' — supported: {', '.join(_SUPPORTED)}")

    if suffix == ".pdf":
        text = _extract_pdf(target)
    elif suffix == ".docx":
        text = _extract_docx(target)
    else:
        text = target.read_text(encoding="utf-8", errors="replace")

    if not text.strip():
        return "(document contains no extractable text — it may be scanned images)"
    return _truncate(text)


@server.tool()
def list_documents(folder: str = "") -> list[str]:
    """List PDF/DOCX files under a vault-relative folder ('' = whole vault)."""
    root = _safe(folder) if folder else VAULT
    out: list[str] = []
    for pattern in ("*.pdf", "*.docx"):
        for f in sorted(root.rglob(pattern)):
            if ".system" not in f.parts:
                out.append(f.relative_to(VAULT).as_posix())
    return sorted(out)[:200]


if __name__ == "__main__":
    server.run()
